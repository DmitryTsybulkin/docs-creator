import smtplib
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart

from docx import Document


def generate_document(title, text):
    document = Document()

    document.add_heading(title, 0)
    document.add_paragraph(text=text)
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    document.add_heading('Heading, level 1', level=1)
    # document.add_page_break()
    name = title + '.docx'

    document.save(name)
    file = open(name, 'rb', encoding='utf-8')
    send_document_to_mail(to="tsydimasik@gmail.com", subject=title, file=file)


def send_document_to_mail(to, subject, file):
    email_address = 'tsydimasik@gmail.com'
    email_password = '68b4b813f76'
    msg = MIMEMultipart()
    msg['From'] = email_address
    msg['To'] = to
    msg['Subject'] = subject

    try:
        part = MIMEApplication(file.read())
        part.add_header('Content-Disposition', 'attachment', filename=file.name)
        msg.attach(part)
    except FileNotFoundError:
        print('Something go wrong')

    server = smtplib.SMTP('smtp.gmail.com', port=587)
    server.ehlo()
    server.starttls()
    server.login(email_address, email_password)
    server.send_message(msg=msg, from_addr=email_address, to_addrs=to)
    return True
